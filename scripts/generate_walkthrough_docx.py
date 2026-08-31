"""
Generates docs/WALKTHROUGH.docx from the artifacts scripts/run_pipeline.py produced, so figures
in the document stay in sync with the trained model rather than being typed in by hand.

Run after scripts/run_pipeline.py has completed.
"""

from __future__ import annotations

import json
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt

ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(ROOT))

MODELS_DIR = ROOT / "models"
REPORTS_DIR = MODELS_DIR / "reports"


def load_json(path: Path) -> dict:
    with open(path) as f:
        return json.load(f)


def h(doc, text, level=1):
    doc.add_heading(text, level=level)


def p(doc, text, bullet=False):
    doc.add_paragraph(text, style="List Bullet" if bullet else None)


def main():
    from src.identify.taxonomy import load_taxonomy

    metrics = load_json(MODELS_DIR / "defense" / "metrics.json")
    tab_fidelity = load_json(REPORTS_DIR / "tabular_fidelity.json")
    graph_fidelity = load_json(REPORTS_DIR / "graph_fidelity.json")
    closed_loop = load_json(REPORTS_DIR / "closed_loop.json")
    taxonomy = load_taxonomy()
    simulated = taxonomy[taxonomy["simulated"]]

    doc = Document()
    doc.add_heading("AI Defense Lab for Payment Security", level=0)
    p(doc, "Solution Walkthrough")

    # ================================================================== 1
    h(doc, "1. The novel fraud attacks identified")
    p(
        doc,
        f"The taxonomy covers {len(taxonomy)} GenAI-powered payment fraud vectors across "
        f"{taxonomy['category'].nunique()} categories, spanning the payment lifecycle from "
        "onboarding through authentication, transaction, settlement/laundering, and dispute. "
        "Each vector is tied to a specific GenAI capability (voice cloning, video deepfakes, "
        "LLM-driven personalization and automation, autonomous agents) and to the transactional "
        "fingerprint it produces once money moves — the signal a payments system can actually "
        f"observe, as opposed to the GenAI artifact itself. {len(simulated)} of the {len(taxonomy)} "
        "vectors are implemented as attack-generation agents; the rest are documented for "
        "completeness because their signal is outside transaction data (document authenticity, "
        "security properties of internal fraud-operations tooling).",
    )
    p(doc, "Categories:")
    for cat in sorted(taxonomy["category"].unique()):
        count = len(taxonomy[taxonomy["category"] == cat])
        p(doc, f"{cat} — {count} vectors", bullet=True)

    p(doc, "The 10 simulated attack types:")
    for _, row in simulated.iterrows():
        p(doc, f"{row['id']} {row['name']} — {row['mechanism']}", bullet=True)

    # ================================================================== 2
    h(doc, "2. How the system generates and simulates these attacks")
    p(
        doc,
        "Generation runs in three layers, each addressing a different limitation of the layer "
        "before it.",
    )

    h(doc, "2.1 Rule-based attack agents", level=2)
    p(
        doc,
        "10 agents (7 behavioral, 3 network) inject labeled fraud incidents onto a real "
        "transaction backbone (the PaySim mobile-money dataset), encoding each attack type's "
        "exact transactional fingerprint: amounts and victim balances are drawn from the real "
        "backbone's own distribution rather than fixed constants, so injected fraud is scaled and "
        "timed against real statistics. Behavioral agents (card-testing bursts, adaptive "
        "structuring, account-takeover velocity spikes, synthetic-identity bust-out, BEC and "
        "romance-scam proxies, a classifier-evasion probe) each emit a self-contained sequence of "
        "transactions for one account. Network agents (mule layering chains, collusive rings, "
        "coordinated fan-in mule bursts) emit transactions across multiple synthetic accounts "
        "forming a specific topology, because no single transaction in a laundering ring is "
        "anomalous on its own — only the shape of who pays whom is.",
    )

    h(doc, "2.2 Tabular WGAN-GP", level=2)
    p(
        doc,
        "A conditional generative adversarial network scales up the 7 behavioral attack types. "
        "Continuous columns (amount, balances) are encoded with mode-specific normalization: a "
        "Bayesian Gaussian mixture identifies each column's modes, and every value is represented "
        "as a one-hot mode indicator plus a scalar offset within that mode, rather than a single "
        "global mean/standard deviation. This matters because transaction amounts are multimodal "
        "(distinct clusters per transaction type), and a single-Gaussian encoding causes a "
        "generator to mode-collapse onto the dominant cluster. The generator and critic are "
        "trained with the Wasserstein loss and a gradient penalty (WGAN-GP) rather than the "
        "standard binary cross-entropy GAN loss, which is markedly less stable on this data — "
        "the critic's gradient vanishes once it becomes confidently accurate, starving the "
        "generator. Fidelity is measured with a discriminative score (a classifier trained to "
        f"separate real from synthetic rows): {tab_fidelity['discriminative_auc']:.3f} AUC "
        "(0.5 = indistinguishable), a correlation-matrix difference of "
        f"{tab_fidelity['correlation_diff']:.3f}, and per-feature KS statistics.",
    )

    h(doc, "2.3 Graph GAN", level=2)
    p(
        doc,
        "A second generative model, built from hand-written graph convolution layers (dense "
        "adjacency operations, no external graph-learning library), targets the 3 network attack "
        "types. The generator produces node features, an existence mask, and an edge-weight "
        "matrix from noise and a ring-type condition, then refines the node features with one "
        "round of message passing over its own proposed adjacency; the discriminator is a graph "
        "classifier that pools node embeddings into a graph-level real/fake score.",
    )
    p(
        doc,
        f"Result: across a systematic sweep of training regularization, this generator converged "
        "to a degenerate solution — either a fully connected or an empty graph — rather than "
        "realistic mule/ring/fan-in topology, within the project's time budget "
        f"(degree-distribution KS {graph_fidelity['degree_ks']:.2f}; synthetic cycle rate "
        f"{graph_fidelity['synth_cycle_rate']:.2f} against a real rate of "
        f"{graph_fidelity['real_cycle_rate']:.2f}). The discriminator trained normally and is "
        "retained as a structural fidelity scorer. Training-data augmentation for network fraud "
        "uses the rule-based network agents at a larger scale instead of generator samples. The "
        "sweep and this conclusion are in scripts/experiment_graph_gan.py.",
    )

    h(doc, "2.4 Red-Team GAN", level=2)
    p(
        doc,
        "A modification to the generator objective used specifically in the closed-loop step "
        "(Section 4): the generator's loss gains a second term rewarding it for being scored as "
        "legitimate by a live snapshot of the defense classifier, not only for fooling the "
        "critic. Because XGBoost has no gradient with respect to its input, this term is computed "
        "against a small differentiable surrogate network distilled to match the classifier's "
        "output probabilities on the same transformed representation the generator produces — "
        "the standard approach for constructing gradient-based attacks against a non-"
        "differentiable model. Re-distilling the surrogate from the current classifier each "
        "iteration is what makes the generator adapt to the defense's current state rather than "
        "a fixed one.",
    )

    # ================================================================== 3
    h(doc, "3. Detection and mitigation model, with efficacy results")
    p(
        doc,
        "The defense is an XGBoost classifier on engineered features: balance-consistency "
        "residuals (expected vs. actual balance change per transaction), trailing 24-hour "
        "transaction count and amount per account (both sender and receiver), and per-account "
        "transaction count to date. It is trained on real legitimate transactions plus rule-based "
        "and GAN-augmented synthetic fraud, with entire incidents and accounts — not individual "
        "rows — held out for testing, so the same account's history cannot appear on both sides "
        "of the split.",
    )
    m = metrics
    p(doc, f"Precision {m['precision']:.3f}, recall {m['recall']:.3f}, F1 {m['f1']:.3f}, "
           f"ROC-AUC {m['roc_auc']:.3f}, PR-AUC {m['pr_auc']:.3f}.")
    p(
        doc,
        f"False-positive rate on legitimate transactions: {m['false_positive_rate_on_legit']:.4f} "
        f"({m['confusion_matrix']['fp']} of {m['n_legit']} legitimate transactions in the held-out "
        "set).",
    )
    p(
        doc,
        "SHAP (TreeExplainer) attributes each flagged transaction to the engineered features that "
        "drove its score, exposed per-transaction in the live application.",
    )
    p(
        doc,
        "Caveat: the held-out set is drawn from the same generative process as the training set "
        "(real legitimate transactions plus rule-based and GAN-augmented synthetic fraud). Freshly "
        "minted synthetic accounts carry an inherently strong signal (no prior transaction "
        "history), which partly accounts for the near-ceiling recall. The closed-loop result in "
        "Section 4, evaluated against fraud generated specifically to evade the model, is the more "
        "informative measure of generalization; production deployment would require validation "
        "against live transaction data and ongoing monitoring for concept drift.",
    )

    # ================================================================== 4
    h(doc, "4. Real-world feasibility in live payment environments")
    h(doc, "Closing the loop", level=2)
    p(
        doc,
        f"One closed-loop iteration: false negatives on the held-out set identified "
        f"'{closed_loop['target_attack_type']}' as the defense's weakest attack type. A Red-Team "
        "GAN generator was trained against a surrogate of the live classifier, conditioned on "
        "that type, and produced a batch of synthetic fraud optimized to evade the model. On a "
        "second, disjoint batch from the same generator, detection rate moved from "
        f"{closed_loop.get('holdout_detection_rate_before', float('nan')):.1%} to "
        f"{closed_loop.get('holdout_detection_rate_after', float('nan')):.1%} after folding the "
        "first batch into training and retraining. Aggregate held-out recall moved from "
        f"{closed_loop['metrics_before']['recall']:.3f} to "
        f"{closed_loop['metrics_after']['recall']:.3f} — already near its ceiling beforehand, "
        "which is why the disjoint adversarial batch is the metric that demonstrates the loop "
        "closing rather than the aggregate number. False-positive rate on legitimate transactions "
        f"moved from {closed_loop['metrics_before']['false_positive_rate_on_legit']:.4f} to "
        f"{closed_loop['metrics_after']['false_positive_rate_on_legit']:.4f} across the same "
        "retraining step. The application exposes this as a live action: mining, retraining, and "
        "re-measurement run on demand and produce a fresh result each time, not a fixed recorded "
        "number. A second, independent path in the same interface lets a user drive the loop "
        "directly: attacks generated on the Generate Attacks page accumulate into a session pool, "
        "which the Closed Loop page retrains a session-scoped classifier on (incident-level "
        "train/holdout split, under 20 seconds) and reports detection on that user's own "
        "held-out incidents before and after — without altering the committed defense model or "
        "the reference figures above.",
    )
    h(doc, "Feasibility", level=2)
    p(
        doc,
        "Every simulated attack is grounded in its transactional fingerprint rather than an "
        "assumption that a payments system can directly perceive a deepfake or a cloned voice — a "
        "production fraud system only has transaction-level signal to act on. The engineered "
        "features (balance-consistency residuals, trailing velocity, account recency) generalize "
        "to any transaction stream carrying account, amount, timestamp, and counterparty fields, "
        "not specifically to PaySim's schema. The closed-loop retraining pattern maps onto a "
        "fraud team's periodic model refresh cycle, with the Red-Team GAN standing in for "
        "adversarial stress-testing that would otherwise be manual or observed only after losses "
        "occur in production. The graph-GAN result in Section 2.3 is reported as a limitation "
        "rather than omitted: network-fraud detection in this system currently depends on "
        "rule-based topology generation for training data, and a converged graph generator is "
        "future work rather than a solved component.",
    )

    out_path = ROOT / "docs" / "WALKTHROUGH.docx"
    doc.save(str(out_path))
    print(f"Saved {out_path}")


if __name__ == "__main__":
    main()
